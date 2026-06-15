"""Generate the AWS Landing Zone Studio technical manual (.docx).

Run:  python tools/build_manual.py
Output: AWS_Landing_Zone_Studio_Technical_Manual.docx (repo root)
"""

from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

INK = RGBColor(0x23, 0x2F, 0x3E)
AMBER = RGBColor(0xC4, 0x74, 0x00)
MUTED = RGBColor(0x6E, 0x7A, 0x8C)
CODE_BG = "F4F6F9"
CODE_INK = RGBColor(0x1A, 0x24, 0x2E)

doc = Document()

# ---------------------------------------------------------------------------
# Base styles
# ---------------------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for i, sz in ((1, 20), (2, 15), (3, 12.5)):
    h = doc.styles[f"Heading {i}"]
    h.font.name = "Calibri"
    h.font.size = Pt(sz)
    h.font.color.rgb = INK if i > 1 else AMBER
    h.font.bold = True


def _shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def h(text, level=1):
    doc.add_heading(text, level=level)


def para(text="", italic=False, bold=False, color=None, size=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic, r.bold = italic, bold
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    return p


def bullets(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        if isinstance(it, tuple):
            r = p.add_run(it[0] + ": ")
            r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    _shade(p, CODE_BG)
    for line in text.split("\n"):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = CODE_INK
        run.add_break()
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ===========================================================================
# Cover page
# ===========================================================================
band = doc.add_paragraph()
band.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = band.add_run("AWS LANDING ZONE STUDIO")
r.font.name = "Consolas"
r.font.size = Pt(11)
r.font.color.rgb = AMBER
r.bold = True

title = doc.add_paragraph()
tr = title.add_run("Technical Manual")
tr.font.size = Pt(34)
tr.bold = True
tr.font.color.rgb = INK
title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
sr = sub.add_run("Interactive multi-account Landing Zone design simulator, "
                 "Well-Architected assessor, IaC generator, and AI advisor")
sr.font.size = Pt(12)
sr.font.color.rgb = MUTED

doc.add_paragraph()
meta = doc.add_table(rows=0, cols=2)
for k, v in (
    ("Document", "Technical Manual / Architecture & Operations Guide"),
    ("Application", "AWS Landing Zone Studio (Streamlit)"),
    ("Repository", "github.com/ajittgosavii/awslz"),
    ("Version", "1.0"),
    ("Date", date.today().isoformat()),
    ("Audience", "Cloud architects, platform engineers, maintainers"),
):
    cells = meta.add_row().cells
    rr = cells[0].paragraphs[0].add_run(k)
    rr.bold = True
    rr.font.size = Pt(10)
    cells[1].paragraphs[0].add_run(v).font.size = Pt(10)

doc.add_paragraph()
disc = para("Scores and costs produced by the tool are transparent planning estimates for "
            "exploring trade-offs — not a substitute for the official AWS Well-Architected Tool, "
            "the AWS Pricing Calculator, or your security/compliance review.",
            italic=True, color=MUTED, size=9.5)
doc.add_page_break()

# ===========================================================================
# 1. Introduction
# ===========================================================================
h("1. Introduction", 1)
para("AWS Landing Zone Studio is a Streamlit web application that helps cloud architects "
     "design, simulate, score, and operationalize AWS multi-account Landing Zones. Unlike "
     "diagramming tools (which visualize existing infrastructure) or deployment tools (which "
     "provision it), the Studio targets the design-and-decide phase: it lets an architect model "
     "account strategies, network patterns, identity and governance choices, and instantly see "
     "the resulting organization structure, guardrails, a six-pillar Well-Architected score, and "
     "a sourced cost estimate — then generate Infrastructure-as-Code and an AI-reviewed report.")

h("1.1 Capability summary", 2)
table(["Area", "What it provides"], [
    ["Design Studio", "Live org/OU tree, network topology, guardrail (SCP) table, 5-dimension scorecard, sourced cost breakdown, maturity level"],
    ["Simulator", "Side-by-side comparison of all account strategies + growth stress-test"],
    ["Playbooks", "Scenario lifecycle simulations: M&A integration (DX+SD-WAN connect, MGN VM migration), replicate/absorb, hybrid network, divestiture, expansion, compliance, scale-out; one-click replication IaC"],
    ["Guide", "In-app user guide: quick start, per-tab how-to, workflows, FAQ"],
    ["CIDR", "IP address planner: allocate supernet→VPC→subnet with a utilization treemap, overlap check, inspect, subdivide"],
    ["Well-Architected", "25 weighted checks across 6 pillars, each mapped to an official best practice; one-click remediations"],
    ["IaC", "Export Terraform / LZA config / Control Tower checklist; reverse-import and score existing IaC"],
    ["Live Estate", "Read-only AWS Organizations scan with confidence-rated signals; Control Tower confirmed via API"],
    ["Drift & History", "Target/actual snapshots over time, scheduled collection, target-vs-actual gap"],
    ["Scenarios", "Durable saved designs, A/B comparison, JSON import/export, shareable links, comments"],
    ["AI Advisor", "Streaming design review, best-practice comparison, roadmap (Claude / OpenAI)"],
], widths=[1.6, 4.6])

h("1.2 Technology stack", 2)
bullets([
    ("Frontend / runtime", "Streamlit (Python), Plotly charts, Graphviz + streamlit-agraph diagrams"),
    ("Domain logic", "Pure-Python modules (dataclasses), unit-testable and UI-agnostic"),
    ("AI", "Anthropic Claude (default, streaming) and OpenAI GPT-4o fallback"),
    ("AWS integration", "boto3 (read-only Organizations / Control Tower / Pricing)"),
    ("Persistence", "Pluggable PostgreSQL or SQLite behind one API"),
    ("Reporting", "fpdf2 (branded PDF), python-docx (this manual)"),
    ("CI/CD", "GitHub Actions — compile + store tests against real Postgres; scheduled drift collector via OIDC"),
])

# ===========================================================================
# 2. Architecture
# ===========================================================================
h("2. Architecture", 1)
para("The application is layered: a thin Streamlit presentation layer (app.py + ui.py) over "
     "stateless, pure-Python domain modules, with side-effecting integrations (AWS, LLM, "
     "database) isolated in their own modules. All domain logic is deterministic and "
     "independently testable; nothing in lz_core/waf/pricing imports Streamlit.")

h("2.1 Layer map", 2)
table(["Layer", "Modules", "Responsibility"], [
    ["Presentation", "app.py, ui.py", "Tabs, charts, chat, theme, login gate, session state"],
    ["Domain", "lz_core.py, waf.py, pricing.py, maturity.py, fixes.py, roadmap.py", "Account math, scoring, cost, WAF assessment, maturity, remediations"],
    ["Generation", "iac.py, iac_import.py, diagrams.py, interactive_diagrams.py, report.py", "IaC export/import, diagrams, PDF report"],
    ["Integration", "live_aws.py, drift_collector.py, llm.py, store.py, sharing.py", "AWS scan, scheduled collection, LLM, persistence, share links"],
], widths=[1.3, 2.7, 2.6])

h("2.2 Request / data flow", 2)
para("On each interaction Streamlit re-runs app.py top-to-bottom. The current design lives in "
     "st.session_state as an LZDesign dataclass. Derived values (scores, cost, accounts, WAF "
     "assessment) are recomputed each run from that single source of truth:")
code("LZDesign (session_state)\n"
     "      │\n"
     "      ├─ score_design()        → 5-dimension scorecard (+ explain_scores)\n"
     "      ├─ estimate_monthly_cost() → cost items + traceable basis\n"
     "      ├─ total_accounts()      → org account math\n"
     "      └─ waf.assess()          → 6 pillars × weighted checks\n"
     "                                   │\n"
     "                                   ├─ fixes.apply_fix()  (one-click remediation)\n"
     "                                   ├─ report.build_pdf_report()\n"
     "                                   └─ llm.stream_completion()  (AI advisor)")

# ===========================================================================
# 3. Module reference
# ===========================================================================
h("3. Module reference", 1)
table(["Module", "Purpose"], [
    ["app.py", "Streamlit UI — all tabs, charts, chat, session wiring, one-click fixes"],
    ["ui.py", "Enterprise theme (CSS), branded chrome, animated login gate (single/multi-user)"],
    ["lz_core.py", "Domain core: account math, transparent scoring, sourced cost model, SCP recommendations, rule-based design suggester"],
    ["pricing.py", "Sourced AWS list prices + stated usage assumptions + optional live Price List API overlay"],
    ["waf.py", "Well-Architected engine — 25 weighted checks, official best-practice mapping (exact/adapted)"],
    ["fixes.py", "One-click remediation engine — maps findings to design mutations"],
    ["maturity.py", "Maturity levels, next-best-action, peer benchmark profiles"],
    ["playbooks.py", "Scenario lifecycle simulations (M&A, divestiture, expansion, compliance, scale-out) with grounded runbooks"],
    ["guide.py", "In-app user guide tab"],
    ["store.py", "Pluggable persistence (PostgreSQL/SQLite): scenarios, snapshots, comments"],
    ["live_aws.py", "Read-only AWS Organizations scanner + Control Tower confirmation + estate→design mapping"],
    ["drift_collector.py", "Headless CLI — scheduled read-only scan → 'actual' drift snapshot"],
    ["iac.py", "IaC exporters: Terraform / LZA / Control Tower; multi-org replication bundle (.zip); hybrid connectivity scaffold (DX + SD-WAN + TGW/Cloud WAN)"],
    ["iac_import.py", "Reverse mode — parse Terraform (python-hcl2 + regex fallback) / LZA / CloudFormation → scored design"],
    ["sharing.py", "Encode/decode a design to a URL-safe shareable token"],
    ["wizard.py", "Guided 3-step onboarding setup flow"],
    ["interactive_diagrams.py", "Clickable streamlit-agraph org graph with node drill-down"],
    ["gv_diagrams.py", "Graphviz builders for org structure + network topology (renamed to free the 'diagrams' name for the icon library)"],
    ["roadmap.py", "Day 0 / 1 / 2 phased implementation plan + Gantt"],
    ["report.py", "Branded PDF design report (fpdf2) with optional AI executive summary"],
    ["llm.py", "Provider layer: Claude (Anthropic SDK, streaming) + OpenAI fallback"],
], widths=[1.7, 4.5])

# ===========================================================================
# 4. Data model
# ===========================================================================
h("4. Data model — LZDesign", 1)
para("A single dataclass (lz_core.LZDesign) captures the entire design. Every computation in "
     "the app is a pure function of this object.")
table(["Field", "Type", "Meaning / domain"], [
    ["org_size", "str", "Startup | SMB | Mid-market | Enterprise"],
    ["compliance", "list", "PCI-DSS, HIPAA, SOC 2, ISO 27001, FedRAMP, APRA CPS 234, GDPR, NIST 800-53, HITRUST, CIS"],
    ["num_teams", "int", "Application teams (1–50)"],
    ["num_workloads", "int", "Workloads / applications (1–60)"],
    ["environments", "list", "dev | test | staging | prod"],
    ["regions", "list", "Active AWS regions"],
    ["account_strategy", "str", "Single / per-environment / per-workload / per-workload-per-environment"],
    ["network_pattern", "str", "Flat peering / TGW hub-and-spoke / Centralized egress + TGW / Cloud WAN"],
    ["identity_model", "str", "IAM users / IAM Identity Center / External IdP federation"],
    ["governance", "str", "Control Tower / LZA / Custom Organizations+SCPs / None"],
    ["centralized_logging", "bool", "Log Archive account enabled"],
    ["security_tooling", "bool", "Org-wide GuardDuty / Security Hub / Config"],
    ["backup_dr", "bool", "Centralized backup & DR"],
], widths=[1.6, 0.7, 4.5])

# ===========================================================================
# 5. Scoring model
# ===========================================================================
h("5. Scoring model", 1)
para("score_design() returns five 0–100 dimension scores. The model is a transparent rubric: "
     "every dimension starts from a base tied to the account strategy and applies named "
     "adjustments. explain_scores() returns the per-factor breakdown shown in the UI, so no "
     "score is opaque.")
table(["Dimension", "Primary drivers"], [
    ["Security & Blast Radius", "Account isolation strategy, security tooling, logging, identity model, governance"],
    ["Scalability", "Account strategy, network pattern (peering penalised, Cloud WAN/TGW rewarded), account vending"],
    ["Operational Simplicity", "Fleet size penalty, managed governance (Control Tower/LZA), identity sprawl, multi-region"],
    ["Cost Efficiency", "Platform overhead as a share of modeled total spend (see §6)"],
    ["Compliance Readiness", "Logging, threat detection, isolation, managed controls, heavy-framework penalties"],
], widths=[2.0, 4.2])
para("Cost Efficiency uses a documented workload-spend proxy (pricing.WORKLOAD_SPEND_PROXY_"
     "PER_MONTH = $2,500/workload) as the denominator of the overhead ratio, so a lean design "
     "scores high and a wasteful one (high overhead, little workload value) scores low.")

# ===========================================================================
# 6. Cost model
# ===========================================================================
h("6. Cost model", 1)
para("estimate_monthly_cost() derives every line item as published AWS list price × a stated "
     "usage assumption — no magic constants. The return value includes a 'basis' list giving the "
     "formula and price source for each item, surfaced in the UI and PDF. Data-transfer/"
     "processing is included, and a per-region price index scales the whole estimate.")
bullets([
    ("Sourced prices", "pricing.LIST_PRICES — each entry carries value, unit, source, note (e.g. NAT gateway $0.045/hr, TGW attachment $0.05/hr, Cloud WAN core edge $0.50/hr)."),
    ("Usage assumptions", "pricing.USAGE — explicit, editable per-account/per-VPC volumes that turn usage-based services (GuardDuty/Config/Security Hub) into monthly figures."),
    ("Live overlay", "pricing.fetch_live_prices() queries the AWS Price List Query API (read-only) for clean hourly SKUs and overrides the list prices for the caller's region; degrades gracefully."),
    ("Region index", "pricing.REGION_PRICE_INDEX scales the estimate (e.g. sa-east-1 ≈ 1.35× us-east-1)."),
])

# ===========================================================================
# 7. Well-Architected engine
# ===========================================================================
h("7. Well-Architected engine", 1)
para("waf.assess() evaluates the design against 25 landing-zone-relevant checks across the six "
     "pillars. Critical practices weigh double in the pillar score. Each check returns PASS / "
     "WARN / FAIL with a finding and remediation, and is enriched with the closest official AWS "
     "Well-Architected best practice plus a match flag — 'exact' or 'adapted' — so results are "
     "never mistaken for an official audit.")
table(["Pillar", "Example checks"], [
    ["Operational Excellence", "Account vending automation, landing zone as code, account-to-team ratio, multi-region standardization"],
    ["Security", "Account isolation, root/baseline controls, centralized IdP, central logging, org-wide threat detection, residency, incident containment"],
    ["Reliability", "Blast radius, HA network topology, centralized backup/DR, multi-location, service quotas"],
    ["Performance Efficiency", "Network pattern fit, region proximity, shared services"],
    ["Cost Optimization", "Cost allocation by account, shared egress, proportionate overhead, sandbox cost controls"],
    ["Sustainability", "Region footprint, shared infrastructure, utilization visibility"],
], widths=[1.9, 4.3])
para("waf.mapping_table() produces the full Studio-check → official-best-practice mapping shown "
     "in the UI and documented for auditors. WAF_DISCLAIMER and WAF_REFERENCE_URL accompany it.")

# ===========================================================================
# 8. Persistence
# ===========================================================================
h("8. Persistence (store.py)", 1)
para("store.py exposes one API (scenarios, snapshots, comments) over a backend chosen at "
     "runtime. ON CONFLICT upserts are shared by both engines; only the parameter placeholder "
     "differs (? for SQLite, %s for PostgreSQL).")
table(["Backend", "Selected when", "Notes"], [
    ["PostgreSQL", "LZ_DATABASE_URL set (app copies it from st.secrets['DATABASE_URL'])", "Durable + shared with the collector; required on Streamlit Cloud"],
    ["SQLite", "Otherwise", "LZ_DB_PATH or ./lz_scenarios.db; great for local/persistent hosts"],
], widths=[1.2, 3.0, 2.0])
para("IMPORTANT — Streamlit Community Cloud: the container filesystem is ephemeral (wiped on "
     "redeploy/sleep). A local SQLite file there is NOT durable for app-side writes. Configure a "
     "managed Postgres (Neon/Supabase) via DATABASE_URL for real durability. The app shows a "
     "warning when running on SQLite.", bold=False)
para("Tables: scenarios(username, name, data_json, updated_at) · snapshots(username, ts, label, "
     "kind, waf_overall, scores_json, data_json) · comments(scenario_key, ts, author, text). "
     "Snapshot timestamps use microsecond precision so rapid saves do not collide on the "
     "(username, ts) primary key.")

# ===========================================================================
# 9. Live scan & drift
# ===========================================================================
h("9. Live estate scan & drift collection", 1)
para("live_aws.scan_organization() performs a strictly read-only scan (List*/Describe* only). "
     "It confirms Control Tower via the real controltower API (status, version, drift, governed "
     "regions), infers account strategy from OU/naming structure, and rates every signal by "
     "confidence (Confirmed via API vs inferred). It maps the estate onto an LZDesign so the "
     "same WAF engine scores the real environment, and honestly lists what it cannot detect.")
para("When no explicit keys are passed, the scanner uses the default boto3 credential chain "
     "(environment, profile, or an EC2/ECS instance role) — this is what the headless collector "
     "relies on.")
h("9.1 Scheduled drift collector", 2)
para("drift_collector.py runs the scan, scores the estate, and writes an 'actual' snapshot to the "
     "shared store. Point it at the same database as the app (LZ_DATABASE_URL or LZ_DB_PATH).")
code("# instance role / env credentials, into a shared Postgres\n"
     "LZ_DATABASE_URL=postgresql://… \\\n"
     "  python drift_collector.py --user operator --region us-east-1 --label nightly")
para("Read-only IAM required: AWSOrganizationsReadOnlyAccess plus "
     "controltower:ListLandingZones / GetLandingZone / ListEnabledControls.")

# ===========================================================================
# 10. IaC export & import
# ===========================================================================
h("10. Infrastructure as Code", 1)
h("10.1 Export (iac.py)", 2)
bullets([
    ("Terraform", "AWS Organizations + OUs + accounts (for_each vending) + real SCP policy documents (deny-leave-org, region restriction, security-service tamper protection, IMDSv2, suspended deny-all)."),
    ("LZA config", "organization / accounts / global YAML fragments."),
    ("Control Tower checklist", "Markdown setup checklist with Account Factory inputs."),
])
h("10.2 Reverse import (iac_import.py)", 2)
para("Upload or paste a Terraform (.tf), LZA (.yaml), or CloudFormation (.yaml/.json) file; the "
     "parser infers an approximate design with a confidence-rated signal table and scores it with "
     "the WAF engine — a local, transparent take on the AWS WA IaC Analyzer. Terraform parsing "
     "prefers python-hcl2 (robust to formatting; quote-tolerant across hcl2 v4/v5+) and falls "
     "back to a dependency-free regex parser; the signal table reports which ran.")

# ===========================================================================
# 11. Deployment
# ===========================================================================
h("11. Deployment", 1)
h("11.1 Run locally", 2)
code("pip install -r requirements.txt\n"
     "cp .streamlit/secrets.toml.example .streamlit/secrets.toml   # add API key(s)\n"
     "streamlit run app.py")
para("Graphviz binary is needed locally for diagrams (winget/brew/apt install graphviz). On "
     "Streamlit Cloud this is handled by packages.txt.")
h("11.2 Streamlit Community Cloud", 2)
bullets([
    "Push to GitHub, create the app from the repo with main file app.py.",
    "Add secrets (App settings → Secrets): API key(s), login key, and DATABASE_URL for durability.",
    "requirements.txt (Python deps) and packages.txt (Graphviz) are picked up automatically.",
])
h("11.3 Secrets reference", 2)
table(["Secret", "Purpose"], [
    ["ANTHROPIC_API_KEY", "Claude (default AI provider). Aliases: ANTHROPIC_KEY, CLAUDE_API_KEY"],
    ["OPENAI_API_KEY", "OpenAI fallback provider"],
    ["APP_PASSWORD", "Single-key login gate (demo key 'awslz' if omitted)"],
    ["[users] table", "Multi-user mode (username = passphrase); per-user scenario workspace"],
    ["DATABASE_URL", "PostgreSQL connection string for durable storage (required on Streamlit Cloud)"],
], widths=[1.9, 4.3])

# ===========================================================================
# 12. CI/CD & automation
# ===========================================================================
h("12. CI/CD & automation", 1)
table(["Workflow", "Trigger", "Does"], [
    [".github/workflows/ci.yml", "push / PR", "Byte-compiles all modules; runs the store test suite against SQLite AND a real PostgreSQL service container"],
    [".github/workflows/drift-collector.yml", "nightly + manual", "OIDC assume-role → read-only scan → 'actual' snapshot (to Postgres, or committed SQLite DB)"],
], widths=[2.3, 1.2, 2.7])
para("CI no-ops cleanly until configured. The drift workflow requires repo secret "
     "AWS_OIDC_ROLE_ARN (and optionally LZ_DATABASE_URL to write straight to Postgres). The store "
     "test suite (tests/test_store.py) covers scenarios CRUD/upsert, snapshot history/ordering, "
     "and comment threads on both backends.")

# ===========================================================================
# 13. Security
# ===========================================================================
h("13. Security considerations", 1)
bullets([
    ("Read-only AWS access", "Live scan and collector use only List*/Describe* calls; the UI warns to use temporary, read-only credentials, which are never stored."),
    ("Least privilege", "OIDC role for the GitHub Action is scoped to the repo and a read-only org/Control-Tower policy — no static keys."),
    ("Authentication", "Login gate supports a single key or a multi-user table; demo mode otherwise."),
    ("Output escaping", "User-supplied content (scenario comments) is HTML-escaped / rendered as literal text to prevent stored XSS in raw-HTML sinks."),
    ("Generated SCPs", "Exported guardrails (deny-leave-org, region restriction, security-service tamper protection, IMDSv2, suspended deny-all) are starting points — always review before applying."),
    ("Secrets", "API keys and DATABASE_URL live in Streamlit secrets / environment, never in code."),
])

# ===========================================================================
# 14. Configuration reference
# ===========================================================================
h("14. Configuration reference (environment variables)", 1)
table(["Variable", "Used by", "Effect"], [
    ["LZ_DATABASE_URL", "store, app, collector", "PostgreSQL connection string; selects the Postgres backend"],
    ["LZ_DB_PATH", "store, collector", "SQLite file path (default ./lz_scenarios.db)"],
    ["LZ_USER", "collector / Action", "Snapshot owner (must match the app user)"],
    ["AWS_REGION / AWS_ACCESS_KEY_ID / AWS_SECRET_ACCESS_KEY / AWS_SESSION_TOKEN", "collector", "Standard AWS credential/region inputs (or use an instance role)"],
], widths=[2.4, 1.5, 2.3])

# ===========================================================================
# 15. Extending the app
# ===========================================================================
h("15. Extending the application", 1)
bullets([
    ("Add a WAF check", "Add a _check(...) to the relevant pillar function in waf.py and a mapping entry in OFFICIAL_BP; optionally wire a one-click fix in fixes.FIX_ACTIONS."),
    ("Add a one-click fix", "Add the check id → {label, changes} to fixes.FIX_ACTIONS; the WAF tab renders the button automatically."),
    ("Tune cost", "Edit pricing.LIST_PRICES (sourced prices) or pricing.USAGE (assumptions); the basis table updates automatically."),
    ("Adjust scoring", "Edit the labelled factors in lz_core._score_with_breakdown; explain_scores() keeps the UI in sync."),
    ("Add a backend", "store.py isolates SQL behind _execute/_ph; add a branch in _connect and reuse the shared queries."),
    ("Add an IaC format", "Add an exporter in iac.py and/or a parser branch in iac_import.parse_iac."),
])

doc.add_paragraph()
para("— End of manual —", italic=True, color=MUTED)

out = "AWS_Landing_Zone_Studio_Technical_Manual.docx"
doc.save(out)
print("WROTE", out)
